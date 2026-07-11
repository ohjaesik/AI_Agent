# app/ingestion/loaders.py

"""파일 확장자별 텍스트 추출 helper.

txt, markdown, pdf, docx 등 업로드 문서에서 RAG 색인 가능한 plain text를 뽑는다.
"""

from __future__ import annotations

from pathlib import Path

from docx import Document
from pypdf import PdfReader

SUPPORTED_EXTENSIONS = {".txt", ".md", ".markdown", ".pdf", ".docx"}


def read_text_file(path: Path) -> str:
    """read_text_file 함수. 파일 확장자별 텍스트 추출 helper. 입력을 검증/변환해 다음 단계가 사용할 값을 반환한다."""
    for encoding in ["utf-8", "utf-8-sig", "cp949"]:
        try:
            return path.read_text(encoding=encoding)
        except UnicodeDecodeError:
            continue

    return path.read_text(errors="ignore")


def read_docx_file(path: Path) -> str:
    """read_docx_file 함수. 파일 확장자별 텍스트 추출 helper. 입력을 검증/변환해 다음 단계가 사용할 값을 반환한다."""
    document = Document(str(path))
    paragraphs = [paragraph.text.strip() for paragraph in document.paragraphs]
    table_texts: list[str] = []

    for table in document.tables:
        for row in table.rows:
            cells = [cell.text.strip() for cell in row.cells]
            row_text = " | ".join(cell for cell in cells if cell)
            if row_text:
                table_texts.append(row_text)

    return "\n".join([*paragraphs, *table_texts]).strip()


def read_pdf_file(path: Path) -> str:
    """read_pdf_file 함수. 파일 확장자별 텍스트 추출 helper. 입력을 검증/변환해 다음 단계가 사용할 값을 반환한다."""
    reader = PdfReader(str(path))
    page_texts = []

    for page_index, page in enumerate(reader.pages, start=1):
        text = page.extract_text() or ""
        text = text.strip()

        if text:
            page_texts.append(f"[page {page_index}]\n{text}")

    return "\n\n".join(page_texts).strip()


def load_document_text(file_path: str | Path) -> str:
    """파일 확장자에 맞는 loader를 선택해 plain text를 추출한다."""
    path = Path(file_path)

    if not path.exists():
        raise FileNotFoundError(f"File not found: {path}")

    if not path.is_file():
        raise ValueError(f"Path is not a file: {path}")

    suffix = path.suffix.lower()

    if suffix not in SUPPORTED_EXTENSIONS:
        raise ValueError(
            f"Unsupported file extension: {suffix}. "
            f"Supported: {', '.join(sorted(SUPPORTED_EXTENSIONS))}"
        )

    if suffix in {".txt", ".md", ".markdown"}:
        return read_text_file(path)

    if suffix == ".docx":
        return read_docx_file(path)

    if suffix == ".pdf":
        return read_pdf_file(path)

    raise ValueError(f"Unsupported file extension: {suffix}")
