# app/tools/docx_generator.py

from __future__ import annotations

from pathlib import Path
from typing import Any

from docx import Document
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor

FONT_NAME = "맑은 고딕"
COLOR_NAVY = "1F2937"
COLOR_BLUE = "2563EB"
COLOR_LIGHT_BLUE = "EFF6FF"
COLOR_GRAY = "F3F4F6"
COLOR_DARK_GRAY = "374151"
COLOR_GREEN = "ECFDF5"
COLOR_WHITE = "FFFFFF"

STATUS_LABELS = {
    "draft": "Draft for Review",
    "reviewed": "Reviewed PoC Planning Report",
    "final": "Final PoC Planning Report",
}

STATUS_NOTES = {
    "draft": "본 문서는 공식 출처와 입력 데이터 기반의 검토용 자동 생성 보고서입니다.",
    "reviewed": "본 문서는 Human Review 기록을 포함한 PoC 기획 검토 보고서입니다.",
    "final": "본 문서는 승인된 검토 기록을 포함한 PoC 기획 최종 보고서입니다.",
}


def report_status(report_data: dict[str, Any]) -> str:
    status = str(report_data.get("status") or "draft").lower()
    return status if status in STATUS_LABELS else "draft"


def report_status_label(report_data: dict[str, Any]) -> str:
    return STATUS_LABELS[report_status(report_data)]


def report_status_note(report_data: dict[str, Any]) -> str:
    return STATUS_NOTES[report_status(report_data)]


def set_cell_text(
    cell,
    text: Any,
    font_size: int = 9,
    bold: bool = False,
    color: str = COLOR_NAVY,
    align: WD_ALIGN_PARAGRAPH = WD_ALIGN_PARAGRAPH.CENTER,
) -> None:
    cell.text = ""
    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
    paragraph = cell.paragraphs[0]
    paragraph.alignment = align
    paragraph.paragraph_format.space_before = Pt(1)
    paragraph.paragraph_format.space_after = Pt(1)
    paragraph.paragraph_format.line_spacing = 1.15
    run = paragraph.add_run("" if text is None else str(text))
    run.bold = bold
    run.font.size = Pt(font_size)
    run.font.name = FONT_NAME
    run.font.color.rgb = RGBColor.from_string(color)
    run._element.rPr.rFonts.set(qn("w:eastAsia"), FONT_NAME)


def set_cell_shading(cell, fill: str = COLOR_GRAY) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shading = OxmlElement("w:shd")
    shading.set(qn("w:fill"), fill)
    tc_pr.append(shading)


def set_cell_border(cell, color: str = "D1D5DB", size: str = "4") -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    borders = tc_pr.first_child_found_in("w:tcBorders")
    if borders is None:
        borders = OxmlElement("w:tcBorders")
        tc_pr.append(borders)

    for edge in ("top", "left", "bottom", "right"):
        tag = f"w:{edge}"
        element = borders.find(qn(tag))
        if element is None:
            element = OxmlElement(tag)
            borders.append(element)
        element.set(qn("w:val"), "single")
        element.set(qn("w:sz"), size)
        element.set(qn("w:space"), "0")
        element.set(qn("w:color"), color)


def set_run_font(run, size: int = 10, bold: bool = False, color: str = COLOR_NAVY) -> None:
    run.font.name = FONT_NAME
    run._element.rPr.rFonts.set(qn("w:eastAsia"), FONT_NAME)
    run.font.size = Pt(size)
    run.font.bold = bold
    run.font.color.rgb = RGBColor.from_string(color)


def add_bottom_border(paragraph, color: str = COLOR_BLUE, size: str = "12") -> None:
    p_pr = paragraph._p.get_or_add_pPr()
    p_bdr = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), size)
    bottom.set(qn("w:space"), "4")
    bottom.set(qn("w:color"), color)
    p_bdr.append(bottom)
    p_pr.append(p_bdr)


def configure_document_styles(doc: Document) -> None:
    section = doc.sections[0]
    section.top_margin = Cm(1.6)
    section.bottom_margin = Cm(1.5)
    section.left_margin = Cm(1.7)
    section.right_margin = Cm(1.7)

    normal = doc.styles["Normal"]
    normal.font.name = FONT_NAME
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), FONT_NAME)
    normal.font.size = Pt(9.5)

    for style_name, size, bold, color in [
        ("Title", 22, True, COLOR_NAVY),
        ("Heading 1", 15, True, COLOR_NAVY),
        ("Heading 2", 12, True, COLOR_DARK_GRAY),
        ("Heading 3", 10.5, True, COLOR_DARK_GRAY),
    ]:
        style = doc.styles[style_name]
        style.font.name = FONT_NAME
        style._element.rPr.rFonts.set(qn("w:eastAsia"), FONT_NAME)
        style.font.size = Pt(size)
        style.font.bold = bold
        style.font.color.rgb = RGBColor.from_string(color)


def add_header_footer(doc: Document, report_data: dict[str, Any]) -> None:
    company_name = report_data.get("company_name") or report_data.get("executive_summary", {}).get("company_name", "")
    status_label = report_status_label(report_data)
    for section in doc.sections:
        header_para = section.header.paragraphs[0]
        header_para.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        header_run = header_para.add_run(f"AX Delivery Planner | {company_name}")
        set_run_font(header_run, size=8, color="6B7280")

        footer_para = section.footer.paragraphs[0]
        footer_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        footer_run = footer_para.add_run(status_label)
        set_run_font(footer_run, size=8, color="9CA3AF")


def add_small_label(doc: Document, text: str) -> None:
    paragraph = doc.add_paragraph()
    paragraph.paragraph_format.space_after = Pt(2)
    run = paragraph.add_run(text)
    set_run_font(run, size=8, bold=True, color=COLOR_BLUE)


def add_paragraph_block(doc: Document, text: str, font_size: int = 9.5) -> None:
    if not text:
        return
    paragraph = doc.add_paragraph()
    paragraph.paragraph_format.line_spacing = 1.38
    paragraph.paragraph_format.space_after = Pt(7)
    paragraph.paragraph_format.first_line_indent = Cm(0.15)
    run = paragraph.add_run(text)
    set_run_font(run, size=font_size, color=COLOR_NAVY)


def add_table_block(
    doc: Document,
    headers: list[Any],
    rows: list[list[Any]],
    font_size: int = 7,
    header_fill: str = COLOR_NAVY,
) -> None:
    if not headers:
        return

    table = doc.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = True

    for idx, header in enumerate(headers):
        cell = table.rows[0].cells[idx]
        set_cell_text(cell, header, font_size=font_size, bold=True, color=COLOR_WHITE)
        set_cell_shading(cell, header_fill)
        set_cell_border(cell, color=header_fill, size="6")

    for row_idx, row in enumerate(rows):
        cells = table.add_row().cells
        fill = COLOR_GRAY if row_idx % 2 == 0 else COLOR_WHITE
        for idx, value in enumerate(row[: len(headers)]):
            align = WD_ALIGN_PARAGRAPH.LEFT if idx in {1, 2, len(headers) - 1} else WD_ALIGN_PARAGRAPH.CENTER
            set_cell_text(cells[idx], value, font_size=font_size, align=align)
            set_cell_shading(cells[idx], fill)
            set_cell_border(cells[idx], color="E5E7EB", size="4")
    doc.add_paragraph()


def add_cover_page(doc: Document, report_data: dict[str, Any]) -> None:
    for _ in range(3):
        doc.add_paragraph()

    label = doc.add_paragraph()
    label.alignment = WD_ALIGN_PARAGRAPH.CENTER
    label_run = label.add_run("AX TRANSFORMATION ASSESSMENT")
    set_run_font(label_run, size=9, bold=True, color=COLOR_BLUE)

    status = doc.add_paragraph()
    status.alignment = WD_ALIGN_PARAGRAPH.CENTER
    status_run = status.add_run(report_status_label(report_data).upper())
    set_run_font(status_run, size=8, bold=True, color="6B7280")

    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title.paragraph_format.space_before = Pt(18)
    title.paragraph_format.space_after = Pt(6)
    title_run = title.add_run(report_data.get("title", "AX 전환 업무 프로세스 진단 보고서"))
    set_run_font(title_run, size=22, bold=True, color=COLOR_NAVY)

    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    subtitle_run = subtitle.add_run(report_data.get("subtitle", "회사 공식자료·RAG·Agent 분석 기반 AI Agent PoC 우선순위 제안"))
    set_run_font(subtitle_run, size=11, color="4B5563")

    line = doc.add_paragraph()
    add_bottom_border(line, color=COLOR_BLUE, size="18")

    for _ in range(5):
        doc.add_paragraph()

    metadata = [
        ("대상 기업", report_data.get("company_name", "")),
        ("최우선 후보", report_data.get("mvp_agent", "")),
        ("작성자", report_data.get("author", "")),
        ("작성일", report_data.get("date", "")),
        ("문서 상태", report_status_label(report_data)),
        ("산출물 성격", "AX PoC 우선순위 진단 / 실행계획 포함"),
    ]
    metadata = [(key, value) for key, value in metadata if value]

    table = doc.add_table(rows=len(metadata), cols=2)
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    for row_idx, (key, value) in enumerate(metadata):
        key_cell = table.rows[row_idx].cells[0]
        value_cell = table.rows[row_idx].cells[1]
        set_cell_text(key_cell, key, font_size=9, bold=True, color=COLOR_WHITE)
        set_cell_shading(key_cell, COLOR_NAVY)
        set_cell_border(key_cell, color=COLOR_NAVY, size="6")
        set_cell_text(value_cell, value, font_size=9, align=WD_ALIGN_PARAGRAPH.LEFT)
        set_cell_shading(value_cell, COLOR_LIGHT_BLUE if row_idx % 2 == 0 else COLOR_WHITE)
        set_cell_border(value_cell, color="E5E7EB", size="4")

    for _ in range(2):
        doc.add_paragraph()

    note = doc.add_paragraph()
    note.alignment = WD_ALIGN_PARAGRAPH.CENTER
    note_run = note.add_run(report_status_note(report_data))
    set_run_font(note_run, size=8, color="6B7280")
    doc.add_page_break()


def add_metric_card(table, row: int, col: int, label: str, value: Any, fill: str = COLOR_LIGHT_BLUE) -> None:
    cell = table.cell(row, col)
    set_cell_shading(cell, fill)
    set_cell_border(cell, color="BFDBFE", size="6")
    cell.text = ""
    p1 = cell.paragraphs[0]
    p1.alignment = WD_ALIGN_PARAGRAPH.LEFT
    p1.paragraph_format.space_after = Pt(2)
    r1 = p1.add_run(label)
    set_run_font(r1, size=7, bold=True, color="6B7280")
    p2 = cell.add_paragraph()
    p2.alignment = WD_ALIGN_PARAGRAPH.LEFT
    p2.paragraph_format.space_after = Pt(0)
    r2 = p2.add_run("" if value is None else str(value))
    set_run_font(r2, size=12, bold=True, color=COLOR_NAVY)


def add_executive_dashboard(doc: Document, report_data: dict[str, Any]) -> None:
    summary = report_data.get("executive_summary") or {}
    candidates = report_data.get("top_candidates") or []
    doc.add_heading("Executive Dashboard", level=1)

    intro = doc.add_paragraph()
    intro.paragraph_format.line_spacing = 1.3
    intro.paragraph_format.space_after = Pt(8)
    intro_run = intro.add_run("분석 결과를 의사결정자가 빠르게 확인할 수 있도록 핵심 지표, 우선 후보, 근거 기반 점수 구조를 요약한다.")
    set_run_font(intro_run, size=9.5, color=COLOR_DARK_GRAY)

    metrics = [
        ("분석 기업", summary.get("company_name")),
        ("분석 업무 수", summary.get("process_count")),
        ("근거 문서 수", summary.get("used_source_count")),
        ("RAG 근거 수", summary.get("evidence_count")),
        ("최우선 Agent", summary.get("top_agent")),
        ("최종 점수", summary.get("top_score")),
        ("월 예상 절감액", summary.get("top_monthly_saving")),
        ("Human Review", summary.get("human_decision")),
    ]

    metric_table = doc.add_table(rows=2, cols=4)
    metric_table.style = "Table Grid"
    metric_table.alignment = WD_TABLE_ALIGNMENT.CENTER
    for idx, (label, value) in enumerate(metrics):
        fill = COLOR_GREEN if label in {"월 예상 절감액", "Human Review"} else COLOR_LIGHT_BLUE
        add_metric_card(metric_table, idx // 4, idx % 4, label, value, fill=fill)

    doc.add_paragraph()
    if candidates:
        add_small_label(doc, "TOP 5 AI AGENT CANDIDATES")
        rows = [[item.get("rank"), item.get("candidate_agent_name"), item.get("final_score"), item.get("discovery_bonus"), item.get("saving_rate"), item.get("monthly_saving"), item.get("status")] for item in candidates]
        add_table_block(doc, ["순위", "후보 Agent", "최종점수", "근거보정", "절감률", "월 절감액", "상태"], rows, font_size=7, header_fill=COLOR_BLUE)
    doc.add_page_break()


def add_table_of_contents(doc: Document, sections: list[dict[str, Any]]) -> None:
    doc.add_heading("Table of Contents", level=1)
    toc_items = ["Executive Dashboard", "Top Candidate Detail"] + [section.get("heading", "제목 없음") for section in sections] + ["References"]
    for idx, heading in enumerate(toc_items, start=1):
        paragraph = doc.add_paragraph()
        paragraph.paragraph_format.line_spacing = 1.25
        paragraph.paragraph_format.space_after = Pt(3)
        run = paragraph.add_run(f"{idx}. {heading}")
        set_run_font(run, size=9.5, color=COLOR_NAVY)
    doc.add_page_break()


def add_top_candidate_details(doc: Document, report_data: dict[str, Any]) -> None:
    candidates = report_data.get("top_candidates") or []
    if not candidates:
        return
    doc.add_heading("Top Candidate Detail", level=1)

    for item in candidates[:5]:
        title = doc.add_paragraph()
        title.paragraph_format.space_before = Pt(8)
        title.paragraph_format.space_after = Pt(3)
        run = title.add_run(f"#{item.get('rank')} {item.get('candidate_agent_name')} — {item.get('process_name')}")
        set_run_font(run, size=11, bold=True, color=COLOR_NAVY)
        add_bottom_border(title, color="DBEAFE", size="8")

        info_table = doc.add_table(rows=2, cols=4)
        info_table.style = "Table Grid"
        info_table.alignment = WD_TABLE_ALIGNMENT.CENTER
        values = [
            ("최종점수", item.get("final_score")),
            ("기본점수", item.get("base_score")),
            ("근거보정", item.get("discovery_bonus")),
            ("월 절감액", item.get("monthly_saving")),
            ("대상 사용자", item.get("target_user")),
            ("상태", item.get("status")),
            ("절감률", item.get("saving_rate")),
            ("위험 플래그", ", ".join(item.get("risk_flags") or []) or "standard_review"),
        ]
        for idx, (label, value) in enumerate(values):
            add_metric_card(info_table, idx // 4, idx % 4, label, value, fill=COLOR_GRAY)

        if item.get("problem"):
            add_small_label(doc, "문제 정의")
            add_paragraph_block(doc, item.get("problem", ""), font_size=9)
        if item.get("suitability_rationale"):
            add_small_label(doc, "AX 적합성 근거")
            add_paragraph_block(doc, item.get("suitability_rationale", ""), font_size=9)
        score_rationale = item.get("score_rationale") or {}
        if score_rationale:
            rows = [[key, value] for key, value in score_rationale.items()]
            add_table_block(doc, ["평가 항목", "점수 근거"], rows, font_size=7, header_fill=COLOR_DARK_GRAY)
    doc.add_page_break()


def add_section(doc: Document, section: dict[str, Any]) -> None:
    doc.add_heading(section.get("heading", "제목 없음"), level=1)
    for block in section.get("blocks", []):
        block_type = block.get("type")
        if block_type == "paragraph":
            add_paragraph_block(doc, block.get("text", ""))
        elif block_type == "table":
            add_table_block(doc, block.get("headers", []), block.get("rows", []), font_size=block.get("font_size", 7), header_fill=COLOR_NAVY)
        elif block_type == "code":
            add_paragraph_block(doc, block.get("text", ""), font_size=8)
        elif block_type == "page_break":
            doc.add_page_break()
        else:
            add_paragraph_block(doc, str(block))


def format_reference(reference: Any, idx: int) -> str:
    if isinstance(reference, str):
        return f"[{idx}] {reference}"

    parts = [f"[{idx}]"]
    for key in ["citation_label", "author_or_org", "source_name"]:
        value = reference.get(key)
        if value:
            parts.append(str(value))
    if reference.get("source_type"):
        parts.append(f"({reference.get('source_type')})")
    if reference.get("published_date"):
        parts.append(f"Published: {reference.get('published_date')}")
    if reference.get("accessed_date"):
        parts.append(f"Accessed: {reference.get('accessed_date')}")
    if reference.get("source_url"):
        parts.append(str(reference.get("source_url")))
    return " ".join(parts)


def add_references(doc: Document, references: list[Any]) -> None:
    doc.add_page_break()
    doc.add_heading("References", level=1)
    if not references:
        paragraph = doc.add_paragraph()
        run = paragraph.add_run("사용된 참고자료가 없습니다.")
        set_run_font(run, size=9)
        return
    for idx, reference in enumerate(references, start=1):
        paragraph = doc.add_paragraph()
        paragraph.paragraph_format.left_indent = Cm(0.2)
        paragraph.paragraph_format.space_after = Pt(4)
        run = paragraph.add_run(format_reference(reference, idx))
        set_run_font(run, size=8.5, color=COLOR_DARK_GRAY)


def generate_docx_report(report_data: dict[str, Any], output_path: str | Path) -> str:
    output_path = Path(output_path)
    output_path.parent.mkdir(parents=True, exist_ok=True)
    doc = Document()
    configure_document_styles(doc)
    add_header_footer(doc, report_data)

    sections = report_data.get("sections", [])
    add_cover_page(doc, report_data)
    add_executive_dashboard(doc, report_data)
    add_table_of_contents(doc, sections)
    add_top_candidate_details(doc, report_data)

    for idx, section in enumerate(sections):
        add_section(doc, section)
        if idx < len(sections) - 1:
            doc.add_page_break()

    add_references(doc, report_data.get("references", []))
    doc.save(output_path)
    return str(output_path)
